from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(30, 30, 30)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

DB = RGBColor(0, 51, 102)
MB = RGBColor(0, 70, 130)
DK = RGBColor(30, 30, 30)
GR = RGBColor(100, 100, 100)
WH = RGBColor(255, 255, 255)

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def set_cell_text(cell, text, bold=False, size=10, color=DK):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color

def add_cell_para(cell, text, size=8, color=DK):
    p = cell.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)

def tbl_borders(table):
    tbl = table._tbl
    pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="003366"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="003366"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="003366"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="003366"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="003366"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="003366"/>'
        '</w:tblBorders>'))

def styled_table(doc, headers, rows, cw=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_text(c, h, bold=True, size=10, color=WH)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(c, '003366')
    for ri, row in enumerate(rows):
        bg = 'EBF0FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            set_cell_text(c, str(val), size=10, color=DK)
            shade(c, bg)
    tbl_borders(t)
    if cw:
        for row_obj in t.rows:
            for i, w in enumerate(cw):
                row_obj.cells[i].width = Cm(w)
    doc.add_paragraph('')
    return t

def H1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.color.rgb = DB
def H2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.color.rgb = MB
def P(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
def B(bold, normal=''):
    p = doc.add_paragraph()
    r = p.add_run(bold); r.bold = True
    if normal: p.add_run(normal)
    p.paragraph_format.space_after = Pt(6)
def BUL(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)


# =============== COVER PAGE ===============
for _ in range(5): doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BUSINESS PLAN'); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = DB

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TalentLens'); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = DB

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI-Powered Talent and Data Services Platform'); r.font.size = Pt(13); r.font.color.rgb = GR

for _ in range(4): doc.add_paragraph('')
for label, val in [('Prepared by:', '[Your Name]'), ('Roll Number:', '[Your Roll Number]'), ('College:', '[Your College Name]'), ('Course:', 'Entrepreneurship Essentials'), ('Date:', 'April 2026')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ' '); r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = DK
    r2 = p.add_run(val); r2.font.size = Pt(12); r2.font.color.rgb = GR
doc.add_page_break()

# =============== TOC ===============
H1('Table of Contents')
for item in ['Chapter 1: Executive Summary', 'Chapter 2: The Business', 'Chapter 3: Market Demand', 'Chapter 4: Competition', 'Chapter 5: Strategy', 'Chapter 6: Resources', 'Chapter 7: Financial Outlay, Financial Closer, and Projected Financials', 'Chapter 8: Risks, Opportunities, Rewards and Sensitivities', 'References']:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(item); r.font.size = Pt(11)
doc.add_page_break()

# =============== CH1: EXECUTIVE SUMMARY ===============
H1('Chapter 1: Executive Summary')
P('The hiring industry, both in India and globally, is going through a period of significant change. Companies are finding it harder and more expensive to hire skilled professionals. At the same time, millions of talented Indians \u2014 engineers, doctors, researchers, writers \u2014 are unable to access quality work opportunities because the systems that connect talent with employers are outdated, biased, and inefficient. This gap is especially visible in the rapidly growing artificial intelligence industry, where companies urgently need domain experts to help train and evaluate their AI models but struggle to find and manage them at scale.')
P('TalentLens is a technology platform designed to solve this problem from both sides. On the company side, we use AI-powered video interviews to assess and verify candidates quickly and objectively. On the talent side, we give skilled professionals a way to showcase their real abilities \u2014 not just their resume or college name \u2014 and connect them to high-quality contract work from companies around the world.')
P('The platform operates through three business lines. The first is an AI interview and assessment system that scores candidates on technical skills, communication, problem-solving, and professionalism. The second is a talent marketplace where companies can hire these verified professionals for projects. The third is a managed service for AI companies, where we assemble and manage teams of domain experts for data labeling, model evaluation, and training data creation.')
P('Our revenue comes from a 30 to 35 percent commission charged on every hour that a contractor works through the platform. This creates a recurring revenue stream that grows as the platform scales. We are targeting an initial investment of approximately Rs 18 lakhs, with break-even expected at around 30 active contractors, which we aim to reach within the first four months of operations.')
P('India is a natural starting point for this business. The country has one of the largest pools of English-speaking skilled workers in the world, with competitive cost structures that are 60 to 70 percent lower than Western markets. The Indian AI industry alone was valued at approximately $14.3 billion in 2026, and the gig economy has crossed 12 million workers, making it the fastest-growing gig workforce globally.')
P('This plan provides a detailed examination of the business model, market opportunity, competitive landscape, strategy, required resources, financial projections, and the risks and opportunities associated with the venture.')
doc.add_page_break()

# =============== CH2: THE BUSINESS ===============
H1('Chapter 2: The Business')

H2('2.1 Company Overview')
styled_table(doc, ['Detail', 'Description'],
    [['Company Name', 'TalentLens Technologies Private Limited'],
     ['Legal Structure', 'Private Limited Company'],
     ['Registered Office', '[City, State]'],
     ['Industry', 'Human Resources Technology (HR Tech) / AI Services'],
     ['Year of Incorporation', '2026'],
     ['Website', '[To be launched]']],
    cw=[5, 12])
P('The core idea behind TalentLens is straightforward: use artificial intelligence to evaluate people based on what they can actually do, not based on where they studied or who they know. Then, use that evaluation to connect them with companies that need their skills. And finally, manage the entire working relationship so that both sides have a smooth, trustworthy experience.')

H2('2.2 Vision and Mission')
B('Vision: ', 'To become the most trusted bridge between India\u2019s skilled workforce and the world\u2019s leading technology companies.')
B('Mission: ', 'To use artificial intelligence to evaluate talent fairly, connect professionals with meaningful work, and help companies build teams faster and better than traditional hiring allows.')

H2('2.3 The Problem We Are Solving')
P('To understand why TalentLens needs to exist, it is important to understand the problems on both sides of the hiring market.')

B('Problem for Companies:')
P('Hiring has become slow, expensive, and unreliable. According to industry reports, the average time to fill a technical role in India is 45 to 60 days. Recruitment agencies typically charge 15 to 25 percent of a candidate\u2019s annual salary as a placement fee. For a position with an annual salary of Rs 12 lakhs, this means spending Rs 1.8 to 3 lakhs just to find one person. And even after this investment, there is no guarantee of quality. HireRight\u2019s Global Benchmark Report has documented that a significant portion of resumes contain exaggerated or false information.')
P('Beyond general hiring challenges, AI companies face an additional and very specific problem. Building and improving AI models requires extensive human involvement. Doctors need to verify whether an AI\u2019s medical advice is accurate. Lawyers need to check legal summaries. Software engineers need to review AI-generated code. These companies need hundreds or thousands of such experts, often on short notice, and managing this workforce is a massive operational challenge.')

B('Problem for Professionals:')
P('India produces millions of graduates every year. A large number are genuinely talented, but they face structural barriers. A skilled developer in a tier-2 or tier-3 city often does not have the professional network to find remote work with global companies. In traditional job applications, a candidate from a lesser-known college is frequently filtered out before anyone evaluates their actual skills. Freelancers face payment delays, scope changes, and disputes. And there is no structured way for skilled professionals to enter the high-paying AI data economy, even though their expertise is exactly what AI companies need.')

H2('2.4 Our Solution')
B('Business Line 1 \u2014 AI Interview and Assessment')
P('When a candidate signs up, they take a 20-minute AI-powered video interview. The system evaluates them on four parameters: technical knowledge (40%), communication clarity (20%), problem-solving approach (25%), and professionalism (15%). They receive a TalentLens Score out of 100, which becomes their verified profile on the platform.')

B('Business Line 2 \u2014 Talent Marketplace')
P('Verified candidates join the talent pool. Companies can search by skill, domain, experience, and score. They hire contractors for 2 weeks to 12 months. TalentLens handles contracts, time tracking, weekly reports, and payments. The company gets one invoice; the contractor gets guaranteed on-time payment.')

B('Business Line 3 \u2014 AI Data Services')
P('We assemble and manage teams of domain experts for AI companies. If a client needs 25 radiologists to label medical images for three months, we handle the entire process \u2014 finding, verifying, onboarding, managing, and delivering.')

H2('2.5 Revenue Model')
P('TalentLens operates on a commission-based model. For every hour a contractor works through the platform, we charge the client a markup of 30 to 35 percent on top of the contractor\u2019s base rate.')
styled_table(doc, ['Component', 'Example'],
    [['Contractor base rate', 'Rs 600 per hour'],
     ['Company pays TalentLens', 'Rs 790 to Rs 810 per hour'],
     ['TalentLens revenue per hour', 'Rs 190 to Rs 210']],
    cw=[7, 10])
P('Revenue is recurring, scales naturally with the talent pool, and does not require selling new products \u2014 growth comes from more contractors working more hours.')

# ===== BMC 9-BLOCK GRID =====
H2('2.6 Business Model Canvas')
P('The Business Model Canvas below summarizes the nine building blocks of TalentLens, following the standard framework developed by Alexander Osterwalder.')

# Create 4-row, 5-column table for BMC grid
bmc = doc.add_table(rows=4, cols=5)
bmc.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_borders(bmc)

# Set all column widths
for row_obj in bmc.rows:
    for i in range(5):
        row_obj.cells[i].width = Cm(3.4)

# --- ROW 0: Headers ---
bmc_headers = ['KEY PARTNERS', 'KEY ACTIVITIES', 'VALUE PROPOSITION', 'CUSTOMER RELATIONSHIPS', 'CUSTOMER SEGMENTS']
for i, h in enumerate(bmc_headers):
    c = bmc.rows[0].cells[i]
    set_cell_text(c, h, bold=True, size=8, color=WH)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade(c, '003366')

# --- ROW 1: Content for KP, KA, VP, CR, CS ---
bmc_row1 = [
    '- Cloud providers (AWS)\n- Payment gateways (Razorpay, Stripe)\n- Engineering and medical colleges\n- Startup incubators',
    '- AI interview development\n- Candidate screening\n- Client acquisition\n- Data labeling project delivery\n- Quality control',
    '- Verified talent in days, not months\n- 60-70% cost saving vs US/Europe\n- AI scoring removes bias\n- End-to-end managed service\n- Expert teams for AI data work',
    '- Dedicated account managers\n- Free pilot program\n- Weekly progress reports\n- Feedback loop after each project',
    '- AI/ML companies for data labeling\n- Tech startups hiring contractors\n- Global companies seeking Indian remote talent',
]
for i, content in enumerate(bmc_row1):
    c = bmc.rows[1].cells[i]
    set_cell_text(c, content, size=8, color=DK)
    shade(c, 'EBF0FA')

# --- ROW 2: Headers for KR and CH (others blank) ---
bmc_row2_headers = ['', 'KEY RESOURCES', '', 'CHANNELS', '']
bmc_row2_content = [
    '',
    '- AI interview scoring engine\n- Verified talent database\n- Project management team\n- Cloud infrastructure (AWS)',
    '',
    '- LinkedIn and email outreach\n- Industry conferences\n- Content marketing\n- Campus drives\n- Referral program',
    '',
]
for i in range(5):
    c = bmc.rows[2].cells[i]
    if bmc_row2_headers[i]:
        set_cell_text(c, bmc_row2_headers[i], bold=True, size=8, color=WH)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(c, '003366')
        add_cell_para(c, bmc_row2_content[i], size=8, color=DK)
    else:
        set_cell_text(c, '', size=8)
        shade(c, 'F5F5F5')

# --- ROW 3: Cost Structure (merge 0-2) and Revenue Streams (merge 3-4) ---
cost_cell = bmc.cell(3, 0).merge(bmc.cell(3, 1)).merge(bmc.cell(3, 2))
rev_cell = bmc.cell(3, 3).merge(bmc.cell(3, 4))

set_cell_text(cost_cell, 'COST STRUCTURE', bold=True, size=8, color=WH)
cost_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
shade(cost_cell, '003366')
add_cell_para(cost_cell, '- Employee salaries (dev, sales, ops)\n- Cloud hosting and AI model APIs\n- Marketing and client acquisition\n- Co-working office rent\n- Payment processing fees\n- Legal and compliance costs', size=8, color=DK)

set_cell_text(rev_cell, 'REVENUE STREAMS', bold=True, size=8, color=WH)
rev_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
shade(rev_cell, '003366')
add_cell_para(rev_cell, '- Primary: 30-35% hourly commission on contractor billing\n- Secondary: AI Data Services project fees\n- Future (Year 2): AI Interview as a Service sold to other companies', size=8, color=DK)

doc.add_paragraph('')
doc.add_page_break()


# =============== CH3: MARKET DEMAND ===============
H1('Chapter 3: Market Demand')

H2('3.1 Industry Overview and Key Trends')
P('The opportunity for TalentLens is shaped by three large industry trends that are creating strong demand for the kind of services we offer.')

B('Trend 1 \u2014 Growing demand for human expertise in AI development')
P('Every time a company like OpenAI, Google, or Meta trains a new AI model, it needs vast amounts of high-quality data that has been reviewed, labeled, and evaluated by humans. A language model needs humans to judge whether its responses are helpful, accurate, and safe. A medical AI needs doctors to verify its diagnoses. As AI models become more sophisticated, the quality bar for this human input goes up, which means companies need more skilled people, not fewer. The Indian AI market alone was valued at approximately $14.3 billion in 2026 (VynZ Research, 2025).')

B('Trend 2 \u2014 Explosive growth of India\u2019s gig economy')
P('India\u2019s gig workforce crossed 12 million workers in 2025 and is projected to reach 23.5 million by 2030, growing at a compound annual growth rate of 21 percent \u2014 the fastest of any country in the world (DemandSage, 2026). Over 15 million freelancers are currently active on various digital platforms. This tells us that both workers and companies are increasingly comfortable with flexible, contract-based work arrangements.')

B('Trend 3 \u2014 Paradoxical talent shortage')
P('Despite having one of the world\u2019s largest graduate populations, 82 percent of employers in India reported difficulty filling positions in 2026, with AI-related skills being the hardest to find (CXO Today, 2026). The India Skills Report 2026 found that India\u2019s employability rate has reached 56.35 percent. The problem is clearly not a lack of talent but a failure of the systems that match talent with opportunity. This is precisely the gap that TalentLens addresses.')

H2('3.2 Target Customer Segments')
styled_table(doc, ['Segment', 'Description', 'Why They Need Us'],
    [['AI and ML Companies', 'Companies building large language models, computer vision, and other AI products', 'Need domain experts at scale for data labeling, evaluation, and training data review'],
     ['Tech Startups and Mid-size Companies', 'Companies needing engineers, data scientists, or designers on contract', 'Want flexibility of contract hiring without slow full-time recruitment'],
     ['Global Companies', 'Companies in US, Europe, Middle East seeking remote Indian professionals', 'Cost advantage: Indian developer Rs 20-30 lakhs/yr vs $150,000-$200,000 in US']],
    cw=[4, 6.5, 6.5])

H2('3.3 Market Size Estimation')
P('India has approximately 1.25 million AI professionals projected by 2027, and the country commands roughly 16 percent of the world\u2019s AI talent (India Skills Report, 2026). Even with a very conservative market capture \u2014 reaching just 500 active contractors in Year 1 (approximately 0.04 percent of the AI professional pool) \u2014 the business can generate meaningful revenue, as detailed in Chapter 7.')
P('The total addressable market for AI data services globally is estimated at several billion dollars, and India, as the world\u2019s largest supplier of data labeling talent, is positioned to capture a significant share.')
P('Note: All market size figures are drawn from published industry reports cited in the References section. Revenue projections are clearly labeled as estimates based on stated assumptions.')
doc.add_page_break()


# =============== CH4: COMPETITION ===============
H1('Chapter 4: Competition')

H2('4.1 Competitive Landscape')
P('The market that TalentLens operates in has several types of existing players, each with their own strengths and limitations.')
styled_table(doc, ['Competitor', 'Type', 'Strengths', 'Weaknesses'],
    [['Upwork / Fiverr', 'Open freelance marketplace', 'Massive user base, global reach', 'No skill verification, pricing race to bottom'],
     ['Toptal', 'Curated freelance network', 'Strong brand, high-quality talent', 'Very expensive, slow onboarding, rejects 97% of applicants'],
     ['Scale AI', 'AI data labeling platform', 'Large scale, major AI lab clients', 'Low-skill crowd workers, not suited for expert-level work'],
     ['Appen', 'Data annotation services', 'Established player, global workforce', 'Quality inconsistency, recent financial difficulties'],
     ['Traditional Agencies', 'Recruitment services', 'Deep industry relationships', 'Slow (45-60 days), manual, 15-25% placement fees']],
    cw=[3.5, 3.5, 5, 5])

H2('4.2 Competitive Positioning')
P('TalentLens occupies a position in the market that none of the existing players fully serve. We combine three capabilities that are currently available only separately:')
BUL('AI-based skill assessment that objectively tests candidates before they are presented to any client')
BUL('A fully managed marketplace that handles the entire lifecycle from matching to payments')
BUL('Specialized AI data services using verified domain experts, not general crowd workers')
P('Our positioning sits in the gap between cheap-but-unreliable platforms (like Upwork) and elite-but-expensive ones (like Toptal). We deliver verified quality at competitive prices, with fast turnaround.')

H2('4.3 SWOT Analysis')
# Proper 2x2 SWOT matrix
swot = doc.add_table(rows=3, cols=3)
swot.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_borders(swot)

# Top-left corner empty
set_cell_text(swot.cell(0,0), '', size=9)
shade(swot.cell(0,0), 'FFFFFF')

# Column headers
set_cell_text(swot.cell(0,1), 'HELPFUL', bold=True, size=9, color=WH)
swot.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
shade(swot.cell(0,1), '003366')

set_cell_text(swot.cell(0,2), 'HARMFUL', bold=True, size=9, color=WH)
swot.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
shade(swot.cell(0,2), '003366')

# Row headers
set_cell_text(swot.cell(1,0), 'INTERNAL', bold=True, size=9, color=WH)
swot.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
shade(swot.cell(1,0), '003366')

set_cell_text(swot.cell(2,0), 'EXTERNAL', bold=True, size=9, color=WH)
swot.cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
shade(swot.cell(2,0), '003366')

# STRENGTHS
set_cell_text(swot.cell(1,1), 'STRENGTHS', bold=True, size=9, color=DB)
shade(swot.cell(1,1), 'E8F5E9')
for s in ['AI-based objective screening', 'Cost-effective Indian talent focus', 'Three diversified revenue lines', 'Low fixed costs (cloud + remote model)']:
    add_cell_para(swot.cell(1,1), '- ' + s, size=8)

# WEAKNESSES
set_cell_text(swot.cell(1,2), 'WEAKNESSES', bold=True, size=9, color=DB)
shade(swot.cell(1,2), 'FFEBEE')
for w in ['New brand, no track record', 'AI technology still developing', 'Small founding team', 'Limited capital vs funded competitors']:
    add_cell_para(swot.cell(1,2), '- ' + w, size=8)

# OPPORTUNITIES
set_cell_text(swot.cell(2,1), 'OPPORTUNITIES', bold=True, size=9, color=DB)
shade(swot.cell(2,1), 'E3F2FD')
for o in ['AI industry growing rapidly', 'India gig economy at 21% CAGR', 'Interview tool as standalone product', 'Government startup support programs']:
    add_cell_para(swot.cell(2,1), '- ' + o, size=8)

# THREATS
set_cell_text(swot.cell(2,2), 'THREATS', bold=True, size=9, color=DB)
shade(swot.cell(2,2), 'FFF3E0')
for t in ['Large platforms adding AI screening', 'Economic slowdowns reducing demand', 'Regulatory changes on gig work', 'AI may reduce labeling need long-term']:
    add_cell_para(swot.cell(2,2), '- ' + t, size=8)

# Column widths for SWOT
for row_obj in swot.rows:
    row_obj.cells[0].width = Cm(3)
    row_obj.cells[1].width = Cm(7)
    row_obj.cells[2].width = Cm(7)

doc.add_paragraph('')
doc.add_page_break()


# =============== CH5: STRATEGY ===============
H1('Chapter 5: Strategy')

H2('5.1 Go-To-Market Strategy')
P('TalentLens is a two-sided platform, which means we need to attract both companies (demand side) and professionals (supply side) simultaneously. Our go-to-market strategy addresses both in parallel.')

B('Demand Side \u2014 Getting Companies:')
styled_table(doc, ['Phase', 'Timeline', 'Activities'],
    [['Phase 1: Direct Outreach', 'Month 1-3', 'Personal outreach to 200-300 CTOs and AI leads via LinkedIn and email. Free pilot: 2 contractors, 2 weeks, no cost.'],
     ['Phase 2: Content & Partnerships', 'Month 3-6', 'Publish case studies and articles on LinkedIn. Partner with 2-3 university AI research labs for credibility.'],
     ['Phase 3: Inbound & Events', 'Month 6-12', 'Attend 3-4 industry conferences (NASSCOM, AI summits). Case studies and testimonials generate inbound leads.']],
    cw=[4, 3, 10])

B('Supply Side \u2014 Getting Talent:')
BUL('Campus drives at engineering, medical, and law colleges \u2014 positioning AI interview as a free career skill assessment')
BUL('Targeted social media campaigns on Instagram and LinkedIn for tech professionals')
BUL('WhatsApp and Telegram communities organized by domain (e.g., "TalentLens Developers")')
BUL('Referral program: Rs 2,000 per successful contractor referral')
BUL('Job board postings on platforms where freelancers actively seek work')

H2('5.2 Pricing Strategy')
styled_table(doc, ['Service', 'Pricing Model', 'Example / Range'],
    [['Talent Marketplace', '30-35% markup on contractor hourly rate', 'Developer at Rs 700/hr base = Rs 910-945/hr to client'],
     ['AI Data Services (Basic)', 'Per hour per worker', 'Rs 300-500 per hour (image tagging, text categorization)'],
     ['AI Data Services (Expert)', 'Per hour per expert', 'Rs 800-2,000 per hour (medical evaluation, legal review)'],
     ['AI Interview as a Service (Year 2)', 'Per interview', 'Rs 200-500 per interview sold to other companies']],
    cw=[4.5, 5, 7.5])

H2('5.3 Growth Roadmap')
styled_table(doc, ['Year', 'Key Milestones'],
    [['Year 1', 'Prove the model: 500 contractors, 15-20 clients, 3-5 case studies, break-even by Month 4-5'],
     ['Year 2', 'Scale: 2,000 contractors, expand to healthcare/legal/finance domains, launch interview tool as product'],
     ['Year 3', 'Raise Series A funding, grow team to 50+, target 10,000 contractors, expand to US/Europe']],
    cw=[3, 14])

H2('5.4 Operations and Quality Control')
P('Day-to-day operations are organized into three parallel streams: talent operations (processing AI interviews, reviewing scores, maintaining the database), client operations (sales, account management, contract administration), and project delivery (managing data labeling projects, coordinating expert teams, quality auditing).')
P('Quality control measures:')
BUL('AI interview scoring calibrated regularly by comparing against human expert evaluations')
BUL('Clients rate every contractor after each engagement; below 3.0/5 rating leads to removal')
BUL('10-15% of all data labeling output is audited by a separate quality reviewer')
BUL('Maximum acceptable error rate for any delivered project: 5 percent')
doc.add_page_break()


# =============== CH6: RESOURCES ===============
H1('Chapter 6: Resources')

H2('6.1 Founding Team')
styled_table(doc, ['Role', 'Responsibilities', 'Background'],
    [['Founder and CEO', 'Business development, client relationships, strategy, fundraising', '[Describe relevant background]'],
     ['Co-Founder and CTO', 'Platform development, AI interview engine, infrastructure', '[Describe relevant background]']],
    cw=[4, 7, 6])

H2('6.2 Team Structure (Year 1)')
styled_table(doc, ['Role', 'Count', 'Monthly Salary (Rs)'],
    [['Full-stack Developer', '1', '60,000'],
     ['Sales and BD Executive', '1', '35,000'],
     ['Talent Operations Coordinator', '1', '25,000'],
     ['Data Labeling Project Manager', '1', '40,000'],
     ['Marketing (Freelance)', '1', '20,000'],
     ['Total', '5', '1,80,000']],
    cw=[7, 3, 7])
P('The team is planned to grow to 10-12 people by end of Year 1 based on client demand.')

H2('6.3 Technology Resources')
styled_table(doc, ['Resource', 'Purpose', 'Estimated Monthly Cost (Rs)'],
    [['AWS Cloud Hosting', 'Platform, storage, computing', '50,000'],
     ['AI/NLP Model APIs', 'Interview engine capabilities', '20,000'],
     ['Razorpay + Stripe', 'Payment processing (India + International)', '30,000 (variable)'],
     ['Communication Tools', 'Slack, Zoom for team and clients', '5,000'],
     ['Analytics Tools', 'Usage tracking and business metrics', '5,000']],
    cw=[5, 6, 6])

H2('6.4 Physical Infrastructure')
P('The core team will operate from a co-working space in [City] at an estimated cost of Rs 40,000 per month. This provides a professional address, meeting rooms, and flexible seating. All contractors work remotely from their own locations.')

H2('6.5 Advisory Board and Partnerships')
P('We plan to assemble an advisory board of three people: one senior professional from the recruitment industry, one technology advisor experienced in AI product development, and one finance mentor for pricing strategy and fundraising guidance.')
P('Key partnerships include engineering and medical colleges for talent drives, startup incubators (NASSCOM 10K Startups, T-Hub) for mentorship and funding, and the AWS Activate program for cloud credits.')

H2('6.6 Intellectual Property')
P('The primary IP assets are the AI interview scoring algorithm (which improves with more data), the verified talent database, and the data labeling quality management workflow. These create a compounding competitive advantage that strengthens as the platform scales.')
doc.add_page_break()


# =============== CH7: FINANCIALS ===============
H1('Chapter 7: Financial Outlay, Financial Closer, and Projected Financials')
P('Note: All financial figures in this chapter are estimates based on the assumptions clearly stated below. Actual results will depend on market conditions, client acquisition speed, and operational execution.')

H2('7.1 Financial Outlay (Startup Costs)')
styled_table(doc, ['Item', 'Estimated Cost (Rs)'],
    [['Company registration, legal, and compliance', '50,000'],
     ['Platform development (MVP)', '3,00,000'],
     ['AI interview engine development', '2,00,000'],
     ['Office setup (co-working deposit + equipment)', '80,000'],
     ['Marketing and launch campaign', '1,50,000'],
     ['Working capital reserve (6 months buffer)', '10,00,000'],
     ['Miscellaneous and contingency', '20,000'],
     ['Total Estimated Investment', '18,00,000']],
    cw=[11, 6])

H2('7.2 Financial Closer (Funding Plan)')
styled_table(doc, ['Source', 'Amount (Rs)', 'Type'],
    [['Founders\u2019 personal savings', '5,00,000', 'Equity (self-funded)'],
     ['Family and friends round', '5,00,000', 'Convertible note or equity'],
     ['Angel investor / Startup incubator', '8,00,000', 'Equity (est. 10-15% stake)'],
     ['Total', '18,00,000', '']],
    cw=[6, 5, 6])
P('The plan avoids bank debt in the initial phase. The commission-based model is designed to generate positive cash flow within a few months, reducing the need for external borrowing.')

H2('7.3 Revenue Projections (Year 1)')
B('Key Assumptions:')
BUL('Operations begin Month 1; first revenue from Month 3 (first 2 months = setup)')
BUL('Average hourly rate charged to company: Rs 792 (including 32% markup)')
BUL('TalentLens earning per contractor per hour: Rs 192 (markup portion)')
BUL('Average hours per contractor per month: 80')

styled_table(doc, ['Month', 'Active Contractors', 'Gross Billing (Rs)', 'TalentLens Revenue (Rs)'],
    [['Month 1-2', '0', '0', '0 (setup phase)'],
     ['Month 3', '20', '12,67,200', '3,07,200'],
     ['Month 4', '40', '25,34,400', '6,14,400'],
     ['Month 5', '65', '41,18,400', '9,98,400'],
     ['Month 6', '100', '63,36,000', '15,36,000'],
     ['Month 7', '150', '95,04,000', '23,04,000'],
     ['Month 8', '200', '1,26,72,000', '30,72,000'],
     ['Month 9', '270', '1,71,07,200', '41,47,200'],
     ['Month 10', '350', '2,21,76,000', '53,76,000'],
     ['Month 11', '420', '2,66,11,200', '64,51,200'],
     ['Month 12', '500', '3,16,80,000', '76,80,000'],
     ['Year 1 Total', '\u2014', '~Rs 13.4 Crores', '~Rs 3.25 Crores']],
    cw=[3, 4, 5, 5])
P('\u201cTalentLens Revenue\u201d represents only the markup portion \u2014 the company\u2019s actual gross income after paying contractors their base rates.')

H2('7.4 Estimated Monthly Expenses (at ~200 active contractors)')
styled_table(doc, ['Expense Category', 'Estimated Monthly Amount (Rs)'],
    [['Staff salaries (5 employees)', '1,80,000'],
     ['Co-working office rent', '40,000'],
     ['Cloud infrastructure (AWS)', '50,000'],
     ['Marketing and sales', '1,00,000'],
     ['Payment processing fees', '30,000'],
     ['Legal, accounting, compliance', '20,000'],
     ['Communication and software tools', '10,000'],
     ['Miscellaneous', '20,000'],
     ['Total Estimated Monthly Expenses', '4,50,000']],
    cw=[10, 7])

H2('7.5 Projected Profit and Loss (Year 1)')
styled_table(doc, ['Item', 'Estimated Amount (Rs)'],
    [['Total Revenue (markup earnings)', '~3,25,00,000'],
     ['Total Operating Expenses (12 months)', '~54,00,000'],
     ['Estimated Profit Before Tax', '~2,71,00,000'],
     ['Estimated Tax (25% for eligible startups)', '~67,75,000'],
     ['Estimated Net Profit After Tax', '~2,03,25,000']],
    cw=[10, 7])
P('Note: The first 4-5 months will be low-revenue or loss-making. Annual profitability is driven by the strong growth in the second half of the year.')

H2('7.6 Break-Even Analysis')
styled_table(doc, ['Metric', 'Value'],
    [['Monthly fixed costs', 'Rs 4,50,000'],
     ['Revenue per active contractor per month', 'Rs 15,360 (Rs 192 x 80 hours)'],
     ['Break-even point', '~30 active contractors'],
     ['Expected to achieve by', 'Month 4 (from first revenue)'],
     ['Total investment recovery', 'Estimated by Month 7-8']],
    cw=[7, 10])

H2('7.7 Key Financial Indicators')
styled_table(doc, ['Indicator', 'Estimated Value'],
    [['Gross Margin', '32%'],
     ['Net Profit Margin (on gross billing)', '~15%'],
     ['Customer Acquisition Cost', 'Rs 25,000 - 40,000 per client'],
     ['Lifetime Value per Client (12 months)', 'Rs 5,00,000+'],
     ['LTV to CAC Ratio', '12:1 to 20:1'],
     ['Payback Period on Initial Investment', '7-8 months']],
    cw=[8, 9])
doc.add_page_break()


# =============== CH8: RISKS, OPPORTUNITIES, REWARDS, SENSITIVITIES ===============
H1('Chapter 8: Risks, Opportunities, Rewards and Sensitivities')

H2('8.1 Risks and Mitigation Strategies')
styled_table(doc, ['Risk', 'Description', 'Mitigation Strategy'],
    [['Technology Risk', 'AI interview system may not assess accurately in early stages', 'Hybrid model: AI screens first, humans review edge cases. Continuous improvement using client feedback.'],
     ['Client Acquisition', 'Enterprise clients slow to adopt new vendors', 'Free pilot program. Start with fast-deciding startups. Build 3-5 case studies in first 6 months.'],
     ['Talent Supply', 'Not enough quality candidates joining the platform', 'Multiple channels: campus drives, social media, referrals, communities. India has millions in the target pool.'],
     ['Cash Flow', 'Gap between paying contractors and receiving client payments', 'Advance deposits from new clients. 6-month working capital reserve. Net-15 terms for regulars.'],
     ['Competition', 'Larger platforms adding similar features', 'Specialize in AI data services niche. Build high client switching costs through relationships.'],
     ['Regulatory', 'Changes in gig worker laws or data privacy rules', 'Legal advisor from day one. Flexible, compliant contracts. Budget for ongoing compliance.']],
    cw=[3, 5.5, 8.5])

H2('8.2 Opportunities')
P('The AI industry globally is in a sustained growth phase, and the demand for human expertise in AI development is expected to continue for years. TalentLens is positioned directly in the path of this demand.')
P('India has been recognized as a global talent powerhouse, with 56.35% employability (India Skills Report, 2026) and approximately 16% of global AI talent. The AI interview engine has standalone commercial value for recruitment agencies, HR departments, and educational institutions \u2014 a significant additional revenue opportunity from Year 2. Government programs like Startup India provide tax benefits and incubation support. Most importantly, every interview conducted, contractor placed, and project delivered generates proprietary data that makes the platform smarter over time \u2014 a compounding advantage difficult for competitors to replicate.')

H2('8.3 Rewards')
B('For founders: ', 'Opportunity to build a scalable technology business in a high-growth sector. With consistent execution, significant valuation potential within 3 years based on HR tech revenue multiples.')
B('For contractors: ', 'Access to higher-paying, structured work. Domain experts (doctors, researchers) can earn substantially more per hour through expert AI evaluation work.')
B('For clients: ', 'Dramatically faster access to pre-verified talent at competitive rates. Days instead of months.')
B('For the economy: ', 'A formal channel for India\u2019s skilled professionals to participate in the global AI economy, contributing to employment, skill development, and foreign exchange earnings.')

H2('8.4 Sensitivity Analysis')
P('To test the robustness of our projections, we modeled four scenarios:')
styled_table(doc, ['Scenario', 'Assumption Change', 'Impact on Year 1'],
    [['Conservative', '300 contractors (not 500) by year end', 'Revenue ~Rs 1.9 Cr. Still profitable. Break-even shifts to Month 5.'],
     ['Aggressive', '800 contractors by year end', 'Revenue ~Rs 5.2 Cr. Higher costs (team expansion) but significantly higher profit.'],
     ['Pricing Pressure', 'Markup drops from 32% to 22%', 'Revenue per hour falls to Rs 132. Break-even rises to ~43 contractors. Viable but needs more volume.'],
     ['Client Concentration', 'Top 2 clients = 60% revenue; one exits', '~30% revenue loss. Mitigated by targeting 15-20 clients; no single client >15-20% of revenue.']],
    cw=[3.5, 5.5, 8])
P('These scenarios demonstrate that TalentLens has a reasonable margin of safety. The business model does not depend on everything going perfectly \u2014 it remains viable even under conservative conditions.')
doc.add_page_break()


# =============== REFERENCES ===============
H1('References')
refs = [
    'Market Research Future (2026). "India AI Recruitment Market Size, Share and Forecast." MRFR Publications.',
    'CXO Today (2026). "Talent Shortages Rise to 82% in India in 2026, as AI Skills Claim Top Spot." CXO Today Media.',
    'India Skills Report (2026). "India Emerges as Global Talent Powerhouse with 56.35% Employability." Wheebox and CII.',
    'Naukri JobSpeak (2026). "IT Hiring Shows Recovery and AI Momentum Continues." Info Edge India.',
    'DemandSage (2026). "Gig Economy Statistics: Growth and Market Size 2026." DemandSage Research.',
    'VynZ Research (2025). "India Artificial Intelligence Market Size and Growth Report, 2035." VynZ Research.',
    'Down to Earth (2026). "Economic Survey 2026: India\u2019s Fast-Growing Gig Economy." CSE.',
    'Adecco India (2026). "India Tech Hiring Projections \u2014 12-15% Growth." Adecco Group India.',
    'HireRight (2024). "Global Benchmark Report on Employment Background Screening." HireRight Inc.',
    'NASSCOM (2025). "Indian Technology Sector: Strategic Review 2025." NASSCOM.',
    'CII and Wheebox (2026). "India Skills Report 2026." Confederation of Indian Industry.',
    'Ken Research (2025). "India Staffing and Recruiting Market: Forecast to 2030." Ken Research.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    r = p.add_run(f'{i}. {ref}'); r.font.size = Pt(10)


# ========== SAVE ==========
doc.save('/home/raethteam/upanshu/b/Business_Plan_TalentLens.docx')
print('DOCX saved.')
