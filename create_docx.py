"""
Convert Business_Plan_TalentLens.md to a styled .docx file.
Reads the markdown, parses headings/paragraphs/tables, and outputs
a properly formatted Word document with Table Grid borders and blue header shading.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.bold = True
    h.font.size = Pt([0, 16, 13, 11][i])


def add_styled_table(doc, rows_data):
    """Add a table with blue header row and borders."""
    if not rows_data or not rows_data[0]:
        return
    ncols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            if i == 0:
                run.bold = True
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()


def parse_table_block(lines):
    """Parse markdown table lines into list of row lists."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # skip separator rows like |:---|:---|
        if re.match(r'^\|[\s:\-|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    return rows


# -- Read the markdown --
with open('Business_Plan_TalentLens.md', 'r') as f:
    md_lines = f.readlines()

# -- Cover page --
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BUSINESS PLAN')
run.bold = True
run.font.size = Pt(24)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TalentLens')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI-Powered Talent and Data Services Platform')
run.italic = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(24)

for text in ['Prepared by: [Your Name]', 'Roll Number: [Your Roll Number]',
             'College: [Your College Name]', 'Course: Entrepreneurship Essentials',
             'Date: April 2026']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

doc.add_page_break()

# -- Parse and render rest of content --
# Skip until after first "---" (cover page stuff we already handled)
i = 0
# Find the TABLE OF CONTENTS line
while i < len(md_lines) and 'TABLE OF CONTENTS' not in md_lines[i]:
    i += 1

in_table = False
table_lines = []

while i < len(md_lines):
    line = md_lines[i].rstrip('\n')
    stripped = line.strip()

    # Page break on ---
    if stripped == '---':
        if in_table:
            add_styled_table(doc, parse_table_block(table_lines))
            in_table = False
            table_lines = []
        doc.add_page_break()
        i += 1
        continue

    # Table detection
    if stripped.startswith('|') and not in_table:
        in_table = True
        table_lines = [line]
        i += 1
        continue
    elif in_table and stripped.startswith('|'):
        table_lines.append(line)
        i += 1
        continue
    elif in_table and not stripped.startswith('|'):
        add_styled_table(doc, parse_table_block(table_lines))
        in_table = False
        table_lines = []
        # fall through to process current line

    # Headings
    if stripped.startswith('CHAPTER ') and stripped == stripped.upper() and len(stripped) > 10:
        doc.add_heading(stripped, level=1)
        i += 1
        continue
    if stripped in ('TABLE OF CONTENTS', 'CONCLUSION', 'REFERENCES'):
        doc.add_heading(stripped, level=1)
        i += 1
        continue

    # Sub-headings (numbered like 2.1, 4.3 etc.)
    if re.match(r'^\d+\.\d+\s', stripped):
        doc.add_heading(stripped, level=2)
        i += 1
        continue

    # Sub-sub-headings (like "The Opener", "Goals and Objectives", etc.)
    if stripped in ('The Opener', 'Goals and Objectives', 'Strategy', 'Resources',
                    'Basic Financials', 'Risk-Opportunities Canvas (Likelihood vs Impact)'):
        doc.add_heading(stripped, level=3)
        i += 1
        continue

    # Months sub-headings in financials
    if stripped.startswith('Months ') and stripped.endswith(':'):
        doc.add_heading(stripped, level=3)
        i += 1
        continue

    # Empty lines
    if not stripped:
        i += 1
        continue

    # Regular paragraph
    if stripped:
        # Clean up markdown bold markers for docx
        text = stripped
        p = doc.add_paragraph()
        # Simple approach: split on ** for bold segments
        parts = text.split('**')
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if idx % 2 == 1:  # odd index = was between ** markers = bold
                run.bold = True

    i += 1

# Flush any remaining table
if in_table:
    add_styled_table(doc, parse_table_block(table_lines))

doc.save('Business_Plan_TalentLens.docx')
print('Done - Business_Plan_TalentLens.docx created')
